"""
Export Service - Generate PDF and DOCX resume files
Handles formatting, styling, and file generation for resume exports
"""

import os
from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.pdfgen import canvas


# ============================================================================
# DOCX Export Functions
# ============================================================================

def export_resume_to_docx(
    resume_content: Dict[str, Any],
    output_path: Optional[str] = None
) -> bytes:
    """
    Export resume to DOCX format with professional styling.
    
    Args:
        resume_content: Dictionary with resume sections:
            - name, email, phone (contact info)
            - summary (professional summary)
            - experience (list of dicts with company, title, duration, description)
            - education (list of dicts with school, degree, field, year)
            - skills (list of skills)
            - achievements (list of achievements)
            - ats_score (optional ATS score)
            - match_percentage (optional match percentage)
        
        output_path: Optional file path to save DOCX file
    
    Returns:
        Bytes of DOCX file content
    """
    
    # Create Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ========== HEADER: Contact Information ==========
    if resume_content.get('name'):
        name_paragraph = doc.add_paragraph(resume_content['name'].upper())
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_style = name_paragraph.runs[0]
        name_style.font.size = Pt(16)
        name_style.font.bold = True
        name_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    # Contact details
    contact_info = []
    if resume_content.get('email'):
        contact_info.append(resume_content['email'])
    if resume_content.get('phone'):
        contact_info.append(resume_content['phone'])
    if resume_content.get('location'):
        contact_info.append(resume_content['location'])
    
    if contact_info:
        contact_paragraph = doc.add_paragraph(' | '.join(contact_info))
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_style = contact_paragraph.runs[0]
        contact_style.font.size = Pt(10)
        contact_style.font.italic = True
    
    # Add ATS score if available
    if resume_content.get('ats_score') is not None:
        score_para = doc.add_paragraph(
            f"ATS Score: {resume_content['ats_score']:.0f}/100 | "
            f"Match: {resume_content.get('match_percentage', 0):.0f}%"
        )
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_style = score_para.runs[0]
        score_style.font.size = Pt(9)
        score_style.font.color.rgb = RGBColor(0, 128, 0)  # Green
    
    doc.add_paragraph()  # Spacing
    
    # ========== PROFESSIONAL SUMMARY ==========
    if resume_content.get('summary'):
        doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        _style_heading(doc.paragraphs[-1])
        doc.add_paragraph(resume_content['summary'])
        doc.add_paragraph()
    
    # ========== SKILLS ==========
    if resume_content.get('skills'):
        doc.add_heading('SKILLS', level=1)
        _style_heading(doc.paragraphs[-1])
        
        # Group skills if they're formatted with commas
        skills_list = resume_content['skills']
        if isinstance(skills_list, str):
            skills_list = [s.strip() for s in skills_list.split(',')]
        elif isinstance(skills_list, list):
            # If it's a list of strings, might need to join and split
            if skills_list and ',' in str(skills_list[0]):
                skills_list = [s.strip() for skill_group in skills_list 
                              for s in skill_group.split(',')]
        
        # Add skills as bullet points
        for skill in skills_list[:15]:  # Limit to top 15 skills
            doc.add_paragraph(skill, style='List Bullet')
        
        doc.add_paragraph()
    
    # ========== EXPERIENCE ==========
    if resume_content.get('experience'):
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        _style_heading(doc.paragraphs[-1])
        
        experiences = resume_content['experience']
        if not isinstance(experiences, list):
            experiences = [experiences]
        
        for exp in experiences:
            # Job title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp.get('title', 'Position')}")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            company_run = title_para.add_run(f" at {exp.get('company', 'Company')}")
            company_run.italic = True
            
            # Duration
            duration_para = doc.add_paragraph(exp.get('duration', ''))
            duration_style = duration_para.runs[0] if duration_para.runs else None
            if duration_style:
                duration_style.font.size = Pt(10)
                duration_style.font.italic = True
                duration_style.font.color.rgb = RGBColor(128, 128, 128)
            
            # Description
            description = exp.get('description', '')
            if isinstance(description, list):
                for item in description:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(description)
            
            doc.add_paragraph()
    
    # ========== EDUCATION ==========
    if resume_content.get('education'):
        doc.add_heading('EDUCATION', level=1)
        _style_heading(doc.paragraphs[-1])
        
        education = resume_content['education']
        if not isinstance(education, list):
            education = [education]
        
        for edu in education:
            # Degree and Field
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(
                f"{edu.get('degree', 'Degree')}"
            )
            degree_run.bold = True
            degree_run.font.size = Pt(11)
            
            if edu.get('field'):
                field_run = degree_para.add_run(f" in {edu['field']}")
                field_run.italic = True
            
            # School
            school_para = doc.add_paragraph(edu.get('school', 'University'))
            school_style = school_para.runs[0] if school_para.runs else None
            if school_style:
                school_style.font.size = Pt(10)
            
            # Year
            if edu.get('year'):
                year_para = doc.add_paragraph(f"Graduated: {edu['year']}")
                year_style = year_para.runs[0] if year_para.runs else None
                if year_style:
                    year_style.font.size = Pt(9)
                    year_style.font.italic = True
            
            doc.add_paragraph()
    
    # ========== ACHIEVEMENTS ==========
    if resume_content.get('achievements'):
        doc.add_heading('ACHIEVEMENTS', level=1)
        _style_heading(doc.paragraphs[-1])
        
        achievements = resume_content['achievements']
        if not isinstance(achievements, list):
            achievements = [achievements]
        
        for achievement in achievements:
            doc.add_paragraph(achievement, style='List Bullet')
        
        doc.add_paragraph()
    
    # ========== FOOTER ==========
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Generated on {datetime.now().strftime('%B %d, %Y')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    # Optionally save to file
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(doc_bytes.getvalue())
    
    return doc_bytes.getvalue()


def _style_heading(paragraph):
    """Style a heading paragraph"""
    for run in paragraph.runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        run.font.size = Pt(12)


# ============================================================================
# PDF Export Functions
# ============================================================================

def export_resume_to_pdf(
    resume_content: Dict[str, Any],
    output_path: Optional[str] = None
) -> bytes:
    """
    Export resume to PDF format with professional styling.
    
    Args:
        resume_content: Dictionary with resume sections (same as DOCX)
        output_path: Optional file path to save PDF file
    
    Returns:
        Bytes of PDF file content
    """
    
    # Create PDF in memory
    pdf_bytes = BytesIO()
    doc = SimpleDocTemplate(
        pdf_bytes,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch,
    )
    
    # Create styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=RGBColor(0, 51, 102),  # Dark blue
        spaceAfter=6,
        alignment=1,  # Center
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=RGBColor(0, 51, 102),
        spaceAfter=8,
        spaceBefore=8,
        fontName='Helvetica-Bold',
        borderBottom=2,
        borderColor=RGBColor(0, 51, 102)
    )
    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=9,
        alignment=1,  # Center
        textColor=RGBColor(64, 64, 64),
        spaceAfter=4
    )
    
    normal_style = styles['Normal']
    normal_style.fontSize = 10
    
    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=4
    )
    
    # Build content
    story = []
    
    # ========== HEADER ==========
    if resume_content.get('name'):
        story.append(Paragraph(
            resume_content['name'].upper(),
            title_style
        ))
    
    # Contact details
    contact_info = []
    if resume_content.get('email'):
        contact_info.append(resume_content['email'])
    if resume_content.get('phone'):
        contact_info.append(resume_content['phone'])
    if resume_content.get('location'):
        contact_info.append(resume_content['location'])
    
    if contact_info:
        story.append(Paragraph(
            ' | '.join(contact_info),
            contact_style
        ))
    
    # ATS score
    if resume_content.get('ats_score') is not None:
        score_text = (
            f"<b>ATS Score: {resume_content['ats_score']:.0f}/100</b> | "
            f"Match: {resume_content.get('match_percentage', 0):.0f}%"
        )
        story.append(Paragraph(score_text, contact_style))
    
    story.append(Spacer(1, 0.15*inch))
    
    # ========== PROFESSIONAL SUMMARY ==========
    if resume_content.get('summary'):
        story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
        story.append(Paragraph(resume_content['summary'], normal_style))
        story.append(Spacer(1, 0.1*inch))
    
    # ========== SKILLS ==========
    if resume_content.get('skills'):
        story.append(Paragraph('SKILLS', heading_style))
        
        skills_list = resume_content['skills']
        if isinstance(skills_list, str):
            skills_list = [s.strip() for s in skills_list.split(',')]
        
        # Create skills in columns
        skills_text = ', '.join(skills_list[:20])
        story.append(Paragraph(skills_text, normal_style))
        story.append(Spacer(1, 0.1*inch))
    
    # ========== EXPERIENCE ==========
    if resume_content.get('experience'):
        story.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))
        
        experiences = resume_content['experience']
        if not isinstance(experiences, list):
            experiences = [experiences]
        
        for exp in experiences:
            # Title and Company
            title_text = (
                f"<b>{exp.get('title', 'Position')}</b> "
                f"<i>at {exp.get('company', 'Company')}</i>"
            )
            story.append(Paragraph(title_text, normal_style))
            
            # Duration
            if exp.get('duration'):
                duration_style = ParagraphStyle(
                    'Duration',
                    parent=normal_style,
                    textColor=RGBColor(128, 128, 128),
                    fontSize=9
                )
                story.append(Paragraph(exp['duration'], duration_style))
            
            # Description
            description = exp.get('description', '')
            if isinstance(description, list):
                for item in description:
                    story.append(Paragraph(f"• {item}", bullet_style))
            else:
                story.append(Paragraph(description, normal_style))
            
            story.append(Spacer(1, 0.08*inch))
    
    # ========== EDUCATION ==========
    if resume_content.get('education'):
        story.append(Paragraph('EDUCATION', heading_style))
        
        education = resume_content['education']
        if not isinstance(education, list):
            education = [education]
        
        for edu in education:
            degree_text = edu.get('degree', 'Degree')
            if edu.get('field'):
                degree_text += f" in {edu['field']}"
            degree_text = f"<b>{degree_text}</b>"
            
            story.append(Paragraph(degree_text, normal_style))
            story.append(Paragraph(edu.get('school', 'University'), normal_style))
            
            if edu.get('year'):
                year_style = ParagraphStyle(
                    'Year',
                    parent=normal_style,
                    fontSize=9,
                    textColor=RGBColor(128, 128, 128)
                )
                story.append(Paragraph(f"Graduated: {edu['year']}", year_style))
            
            story.append(Spacer(1, 0.08*inch))
    
    # ========== ACHIEVEMENTS ==========
    if resume_content.get('achievements'):
        story.append(Paragraph('ACHIEVEMENTS', heading_style))
        
        achievements = resume_content['achievements']
        if not isinstance(achievements, list):
            achievements = [achievements]
        
        for achievement in achievements:
            story.append(Paragraph(f"• {achievement}", bullet_style))
        
        story.append(Spacer(1, 0.1*inch))
    
    # ========== FOOTER ==========
    footer_style = ParagraphStyle(
        'Footer',
        parent=normal_style,
        fontSize=8,
        textColor=RGBColor(128, 128, 128),
        alignment=1
    )
    story.append(Paragraph(
        f"Generated on {datetime.now().strftime('%B %d, %Y')}",
        footer_style
    ))
    
    # Build PDF
    doc.build(story)
    pdf_bytes.seek(0)
    
    # Optionally save to file
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes.getvalue())
    
    return pdf_bytes.getvalue()


# ============================================================================
# Plain Text Export
# ============================================================================

def export_resume_to_text(
    resume_content: Dict[str, Any],
    output_path: Optional[str] = None
) -> str:
    """
    Export resume to plain text format (ATS-friendly).
    
    Args:
        resume_content: Dictionary with resume sections
        output_path: Optional file path to save text file
    
    Returns:
        Plain text resume content
    """
    
    lines = []
    
    # Header
    if resume_content.get('name'):
        lines.append(resume_content['name'].upper())
        lines.append('=' * len(resume_content['name']))
    
    # Contact
    contact = []
    if resume_content.get('email'):
        contact.append(resume_content['email'])
    if resume_content.get('phone'):
        contact.append(resume_content['phone'])
    if resume_content.get('location'):
        contact.append(resume_content['location'])
    
    if contact:
        lines.append(' | '.join(contact))
    
    if resume_content.get('ats_score') is not None:
        lines.append(f"ATS Score: {resume_content['ats_score']:.0f}/100 | "
                    f"Match: {resume_content.get('match_percentage', 0):.0f}%")
    
    lines.append('')
    
    # Professional Summary
    if resume_content.get('summary'):
        lines.append('PROFESSIONAL SUMMARY')
        lines.append('-' * 20)
        lines.append(resume_content['summary'])
        lines.append('')
    
    # Skills
    if resume_content.get('skills'):
        lines.append('SKILLS')
        lines.append('-' * 20)
        skills = resume_content['skills']
        if isinstance(skills, str):
            lines.append(skills)
        else:
            for skill in skills:
                lines.append(f"• {skill}")
        lines.append('')
    
    # Experience
    if resume_content.get('experience'):
        lines.append('PROFESSIONAL EXPERIENCE')
        lines.append('-' * 20)
        
        experiences = resume_content['experience']
        if not isinstance(experiences, list):
            experiences = [experiences]
        
        for exp in experiences:
            lines.append(f"{exp.get('title')} at {exp.get('company')}")
            if exp.get('duration'):
                lines.append(f"  {exp['duration']}")
            
            description = exp.get('description', '')
            if isinstance(description, list):
                for item in description:
                    lines.append(f"  • {item}")
            else:
                lines.append(f"  {description}")
            lines.append('')
    
    # Education
    if resume_content.get('education'):
        lines.append('EDUCATION')
        lines.append('-' * 20)
        
        education = resume_content['education']
        if not isinstance(education, list):
            education = [education]
        
        for edu in education:
            degree = edu.get('degree', '')
            if edu.get('field'):
                degree += f" in {edu['field']}"
            lines.append(degree)
            lines.append(f"  {edu.get('school')}")
            if edu.get('year'):
                lines.append(f"  Graduated: {edu['year']}")
            lines.append('')
    
    # Achievements
    if resume_content.get('achievements'):
        lines.append('ACHIEVEMENTS')
        lines.append('-' * 20)
        
        achievements = resume_content['achievements']
        if not isinstance(achievements, list):
            achievements = [achievements]
        
        for achievement in achievements:
            lines.append(f"• {achievement}")
        lines.append('')
    
    # Footer
    lines.append(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    
    text_content = '\n'.join(lines)
    
    # Save to file if path provided
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
    
    return text_content


# ============================================================================
# Helper Functions
# ============================================================================

def create_export_dir(base_dir: str = "data/exports") -> str:
    """Create exports directory if it doesn't exist"""
    export_dir = Path(base_dir)
    export_dir.mkdir(parents=True, exist_ok=True)
    return str(export_dir)


def generate_filename(name: str, extension: str) -> str:
    """Generate a unique filename for export"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_name = safe_name.replace(' ', '_')[:30]
    return f"{safe_name}_{timestamp}.{extension}"


# ============================================================================
# Main Export Function
# ============================================================================

def export_resume(
    resume_content: Dict[str, Any],
    export_format: str = "all",
    output_dir: Optional[str] = None
) -> Dict[str, Any]:
    """
    Export resume in multiple formats.
    
    Args:
        resume_content: Dictionary with resume sections
        export_format: 'pdf', 'docx', 'text', or 'all'
        output_dir: Directory to save files (optional)
    
    Returns:
        Dictionary with file bytes and metadata
    """
    
    result = {
        "status": "success",
        "formats": {},
        "timestamp": datetime.now().isoformat()
    }
    
    name = resume_content.get('name', 'resume')
    
    if export_format in ('pdf', 'all'):
        try:
            pdf_bytes = export_resume_to_pdf(resume_content)
            result['formats']['pdf'] = {
                "size_bytes": len(pdf_bytes),
                "generated": True
            }
            
            if output_dir:
                export_dir = create_export_dir(output_dir)
                filename = generate_filename(name, 'pdf')
                filepath = Path(export_dir) / filename
                with open(filepath, 'wb') as f:
                    f.write(pdf_bytes)
                result['formats']['pdf']['filepath'] = str(filepath)
        except Exception as e:
            result['formats']['pdf'] = {"error": str(e), "generated": False}
    
    if export_format in ('docx', 'all'):
        try:
            docx_bytes = export_resume_to_docx(resume_content)
            result['formats']['docx'] = {
                "size_bytes": len(docx_bytes),
                "generated": True
            }
            
            if output_dir:
                export_dir = create_export_dir(output_dir)
                filename = generate_filename(name, 'docx')
                filepath = Path(export_dir) / filename
                with open(filepath, 'wb') as f:
                    f.write(docx_bytes)
                result['formats']['docx']['filepath'] = str(filepath)
        except Exception as e:
            result['formats']['docx'] = {"error": str(e), "generated": False}
    
    if export_format in ('text', 'all'):
        try:
            text_content = export_resume_to_text(resume_content)
            result['formats']['text'] = {
                "size_bytes": len(text_content.encode('utf-8')),
                "generated": True
            }
            
            if output_dir:
                export_dir = create_export_dir(output_dir)
                filename = generate_filename(name, 'txt')
                filepath = Path(export_dir) / filename
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                result['formats']['text']['filepath'] = str(filepath)
        except Exception as e:
            result['formats']['text'] = {"error": str(e), "generated": False}
    
    return result
